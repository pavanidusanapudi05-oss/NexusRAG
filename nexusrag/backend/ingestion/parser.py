import os
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field

@dataclass
class ParsedSection:
    title: str
    content: str
    page_number: int
    section_index: int
    metadata: Dict[str, Any] = field(default_factory=dict)

@dataclass
class ParsedDocument:
    doc_name: str
    file_path: str
    file_type: str
    checksum: str
    metadata: Dict[str, Any]
    sections: List[ParsedSection]
    full_text: str
    total_pages: int

class DocumentParser:
    @staticmethod
    def parse_pdf(file_path: str, metadata: Optional[Dict[str, Any]] = None) -> ParsedDocument:
        import pymupdf # PyMuPDF
        
        doc = pymupdf.open(file_path)
        total_pages = len(doc)
        sections: List[ParsedSection] = []
        all_text_parts = []
        sec_idx = 0

        # Section regex patterns: 'Section 1. ...', 'Section 1: ...', 'Article 2', etc.
        sec_pattern = re.compile(r'^(Section\s+\d+[^\n]*|Article\s+\d+[^\n]*|[A-Z0-9\.\s]{4,40}:)', re.IGNORECASE)

        for page_idx, page in enumerate(doc):
            page_num = page_idx + 1
            page_text = page.get_text('text')
            all_text_parts.append(page_text)
            
            lines = page_text.splitlines()
            current_sec_title = f'Page {page_num} Overview'
            current_sec_lines = []
            
            for line in lines:
                stripped = line.strip()
                if not stripped:
                    continue
                if sec_pattern.match(stripped) and len(stripped) < 80:
                    if current_sec_lines:
                        sections.append(ParsedSection(
                            title=current_sec_title,
                            content='\n'.join(current_sec_lines),
                            page_number=page_num,
                            section_index=sec_idx,
                            metadata={'source_page': page_num}
                        ))
                        sec_idx += 1
                        current_sec_lines = []
                    current_sec_title = stripped
                else:
                    current_sec_lines.append(stripped)
            
            if current_sec_lines:
                sections.append(ParsedSection(
                    title=current_sec_title,
                    content='\n'.join(current_sec_lines),
                    page_number=page_num,
                    section_index=sec_idx,
                    metadata={'source_page': page_num}
                ))
                sec_idx += 1

        doc.close()
        full_text = '\n\n'.join(all_text_parts)
        meta = metadata or {}
        meta.setdefault('total_pages', total_pages)

        return ParsedDocument(
            doc_name=Path(file_path).name,
            file_path=file_path,
            file_type='pdf',
            checksum='',
            metadata=meta,
            sections=sections,
            full_text=full_text,
            total_pages=total_pages
        )

    @staticmethod
    def parse_docx(file_path: str, metadata: Optional[Dict[str, Any]] = None) -> ParsedDocument:
        from docx import Document
        doc = Document(file_path)
        sections: List[ParsedSection] = []
        all_text_parts = []
        
        current_title = 'Overview'
        current_lines = []
        sec_idx = 0
        
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            all_text_parts.append(text)
            
            if p.style.name.startswith('Heading') or text.lower().startswith('section'):
                if current_lines:
                    sections.append(ParsedSection(
                        title=current_title,
                        content='\n'.join(current_lines),
                        page_number=1,
                        section_index=sec_idx,
                        metadata={'style': 'paragraph'}
                    ))
                    sec_idx += 1
                    current_lines = []
                current_title = text
            else:
                current_lines.append(text)
                
        # Parse tables in docx
        for t_idx, table in enumerate(doc.tables):
            table_lines = []
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                table_lines.append(row_text)
            if table_lines:
                table_content = '\n'.join(table_lines)
                all_text_parts.append(table_content)
                sections.append(ParsedSection(
                    title=f'Table {t_idx+1}',
                    content=table_content,
                    page_number=1,
                    section_index=sec_idx,
                    metadata={'is_table': True}
                ))
                sec_idx += 1

        if current_lines:
            sections.append(ParsedSection(
                title=current_title,
                content='\n'.join(current_lines),
                page_number=1,
                section_index=sec_idx,
                metadata={'style': 'paragraph'}
            ))

        meta = metadata or {}
        meta.setdefault('total_pages', 1)

        return ParsedDocument(
            doc_name=Path(file_path).name,
            file_path=file_path,
            file_type='docx',
            checksum='',
            metadata=meta,
            sections=sections,
            full_text='\n\n'.join(all_text_parts),
            total_pages=1
        )

    @staticmethod
    def parse_spreadsheet(file_path: str, metadata: Optional[Dict[str, Any]] = None) -> ParsedDocument:
        import pandas as pd
        
        path = Path(file_path)
        if path.suffix.lower() == '.csv':
            df_dict = {'Sheet1': pd.read_csv(file_path)}
        else:
            df_dict = pd.read_excel(file_path, sheet_name=None)
            
        sections: List[ParsedSection] = []
        all_text_parts = []
        sec_idx = 0
        
        for sheet_name, df in df_dict.items():
            df = df.fillna('')
            sheet_title = f'Sheet: {sheet_name}'
            rows_summary = []
            
            # Format row by row for clear semantic indexing
            for r_idx, row in df.iterrows():
                row_items = [f'{col}: {val}' for col, val in row.items() if str(val).strip()]
                row_str = ' | '.join(row_items)
                rows_summary.append(f'[Row {r_idx+1}] {row_str}')
                
            sheet_content = f'{sheet_title}\nTotal Records: {len(df)}\n' + '\n'.join(rows_summary)
            all_text_parts.append(sheet_content)
            
            sections.append(ParsedSection(
                title=sheet_title,
                content=sheet_content,
                page_number=1,
                section_index=sec_idx,
                metadata={'sheet_name': sheet_name, 'records': len(df)}
            ))
            sec_idx += 1

        meta = metadata or {}
        meta.setdefault('total_pages', len(df_dict))

        return ParsedDocument(
            doc_name=path.name,
            file_path=file_path,
            file_type=path.suffix.lstrip('.').lower(),
            checksum='',
            metadata=meta,
            sections=sections,
            full_text='\n\n'.join(all_text_parts),
            total_pages=len(df_dict)
        )

    @staticmethod
    def parse_text(file_path: str, metadata: Optional[Dict[str, Any]] = None) -> ParsedDocument:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            raw_text = f.read()
            
        lines = raw_text.splitlines()
        sections: List[ParsedSection] = []
        sec_idx = 0
        current_title = 'Section 1'
        current_lines = []
        
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith('#') or stripped.lower().startswith('section'):
                if current_lines:
                    sections.append(ParsedSection(
                        title=current_title,
                        content='\n'.join(current_lines),
                        page_number=1,
                        section_index=sec_idx,
                        metadata={}
                    ))
                    sec_idx += 1
                    current_lines = []
                current_title = stripped.lstrip('#').strip()
            else:
                current_lines.append(stripped)
                
        if current_lines:
            sections.append(ParsedSection(
                title=current_title,
                content='\n'.join(current_lines),
                page_number=1,
                section_index=sec_idx,
                metadata={}
            ))
            
        meta = metadata or {}
        meta.setdefault('total_pages', 1)

        return ParsedDocument(
            doc_name=Path(file_path).name,
            file_path=file_path,
            file_type='txt',
            checksum='',
            metadata=meta,
            sections=sections,
            full_text=raw_text,
            total_pages=1
        )

    @classmethod
    def parse(cls, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> ParsedDocument:
        ext = Path(file_path).suffix.lower()
        if ext == '.pdf':
            return cls.parse_pdf(file_path, metadata)
        elif ext in ['.docx', '.doc']:
            return cls.parse_docx(file_path, metadata)
        elif ext in ['.xlsx', '.xls', '.csv']:
            return cls.parse_spreadsheet(file_path, metadata)
        else:
            return cls.parse_text(file_path, metadata)
