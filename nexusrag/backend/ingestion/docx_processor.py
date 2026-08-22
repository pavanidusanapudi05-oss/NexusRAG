import os
from pathlib import Path
from typing import List, Dict, Any
from docx import Document
from .pdf_processor import ExtractedContentBlock

class DOCXProcessor:
    @staticmethod
    def process(file_path: Path) -> List[ExtractedContentBlock]:
        doc = Document(str(file_path))
        blocks: List[ExtractedContentBlock] = []

        current_heading = "Overview"
        current_lines = []

        # 1. Paragraphs and Headings
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            if p.style.name.startswith("Heading") or text.lower().startswith("section"):
                if current_lines:
                    blocks.append(ExtractedContentBlock(
                        text="\n".join(current_lines),
                        page_number=1,
                        section_title=current_heading,
                        metadata={"style": "paragraph"}
                    ))
                    current_lines = []
                current_heading = text
            else:
                current_lines.append(text)

        if current_lines:
            blocks.append(ExtractedContentBlock(
                text="\n".join(current_lines),
                page_number=1,
                section_title=current_heading,
                metadata={"style": "paragraph"}
            ))

        # 2. Tables
        for t_idx, table in enumerate(doc.tables):
            table_lines = []
            for row in table.rows:
                row_str = " | ".join(cell.text.strip() for cell in row.cells)
                if row_str.strip():
                    table_lines.append(row_str)
            if table_lines:
                blocks.append(ExtractedContentBlock(
                    text="\n".join(table_lines),
                    page_number=1,
                    section_title=f"Table {t_idx+1}",
                    metadata={"is_table": True}
                ))

        return blocks
