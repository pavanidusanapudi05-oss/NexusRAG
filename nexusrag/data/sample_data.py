import os
from pathlib import Path
from typing import List, Optional
import openpyxl
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet

def generate_sample_documents(target_dir: Optional[str] = None) -> List[str]:
    if target_dir is None:
        target_dir = os.path.join(os.path.dirname(__file__), "documents")
    
    os.makedirs(target_dir, exist_ok=True)
    created_files = []

    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    h1_style = styles["Heading1"]
    body_style = styles["BodyText"]

    # 1. Employee Policy 2025 PDF
    doc2025_pdf = os.path.join(target_dir, "Employee_Operations_Policy_2025.pdf")
    if not os.path.exists(doc2025_pdf):
        doc = SimpleDocTemplate(doc2025_pdf, pagesize=letter)
        story = [
            Paragraph("NexusCorp Employee Operations Policy 2025", title_style),
            Paragraph("Document ID: POL-2025-V1 | Version: 1.0 | Effective Year: 2025 | Department: Human Resources & Operations", body_style),
            Spacer(1, 15),
            Paragraph("Section 1. Purpose & Scope", h1_style),
            Paragraph("This policy governs operational workflows, workplace attendance, remote working arrangements, and procurement standards for all full-time and contract personnel across NexusCorp enterprise operations.", body_style),
            Spacer(1, 10),
            Paragraph("Section 2. Working Hours & Attendance", h1_style),
            Paragraph("Standard working hours are 40 hours per week, scheduled Monday through Friday from 9:00 AM to 5:00 PM local time. All employees are required to maintain a minimum of 75% on-site office attendance calculated on a rolling monthly basis. If attendance falls below 75%, an automated warning is issued, requiring immediate review with the HR Department.", body_style),
            Spacer(1, 10),
            Paragraph("Section 3. Remote Work Authorization", h1_style),
            Paragraph("Eligible employees may request a maximum of 1 day per week of remote work. All remote work requests require written approval from the respective Department Manager at least 5 business days in advance. Remote work is subject to maintaining positive performance reviews.", body_style),
            PageBreak(),
            Paragraph("Section 4. Leave Entitlement & Notice Period", h1_style),
            Paragraph("Full-time employees receive 20 days of annual paid leave per calendar year. Leave requests exceeding 3 consecutive business days must be submitted with at least 14 days advance notice. Unused annual leave cannot exceed 5 carryover days into the subsequent calendar year.", body_style),
            Spacer(1, 10),
            Paragraph("Section 5. Travel & Meal Allowances", h1_style),
            Paragraph("Domestic business travel is capped at a daily meal allowance of $50 per day. All domestic flights must be booked in economy class. Lodging expenses require prior budget authorization and are capped at standard city per-diem rates.", body_style),
            Spacer(1, 10),
            Paragraph("Section 6. IT Equipment & Home Office", h1_style),
            Paragraph("NexusCorp provides standard corporate laptop hardware. No additional home office subsidy or peripheral reimbursement is authorized under the 2025 policy.", body_style),
        ]
        doc.build(story)
    created_files.append(doc2025_pdf)

    # 2. Employee Policy 2026 PDF
    doc2026_pdf = os.path.join(target_dir, "Employee_Operations_Policy_2026.pdf")
    if not os.path.exists(doc2026_pdf):
        doc = SimpleDocTemplate(doc2026_pdf, pagesize=letter)
        story = [
            Paragraph("NexusCorp Employee Operations Policy 2026", title_style),
            Paragraph("Document ID: POL-2026-V2 | Version: 2.0 | Effective Year: 2026 | Department: Human Resources & Operations", body_style),
            Spacer(1, 15),
            Paragraph("Section 1. Purpose & Scope", h1_style),
            Paragraph("This policy supersedes Policy POL-2025-V1 and governs modern workplace flexibility, hybrid operations, digital security requirements, and updated reimbursement standards across NexusCorp.", body_style),
            Spacer(1, 10),
            Paragraph("Section 2. Working Hours & Attendance", h1_style),
            Paragraph("NexusCorp adopts flexible working hours totaling 40 hours per week, with mandatory core collaboration hours between 10:00 AM and 3:00 PM. Mandatory on-site office attendance is updated to a minimum of 60% calculated monthly. If attendance falls below 60%, the employee enters a collaborative review process with their Department Manager and HR.", body_style),
            Spacer(1, 10),
            Paragraph("Section 3. Remote Work Authorization", h1_style),
            Paragraph("Eligible employees can work remotely for up to 3 days per week. Remote work authorization requires written approval from the Department Director and compliance with Enterprise Security Regulation SR-402 (MFA and secure VPN).", body_style),
            PageBreak(),
            Paragraph("Section 4. Leave Entitlement & Notice Period", h1_style),
            Paragraph("Full-time employees are entitled to 25 days of annual paid leave per calendar year. Leave requests exceeding 3 consecutive business days now require only 7 days advance notice. Up to 10 days of unused leave may be carried forward into the next year.", body_style),
            Spacer(1, 10),
            Paragraph("Section 5. Travel & Meal Allowances", h1_style),
            Paragraph("Daily meal allowance for domestic travel is increased to $75 per day. Flights exceeding 6 hours in total duration are eligible for premium economy class seating. Incidentals up to $25 per trip may be claimed without itemized receipts.", body_style),
            Spacer(1, 10),
            Paragraph("Section 6. IT Equipment & Home Office", h1_style),
            Paragraph("NexusCorp provides corporate laptops plus a one-time $500 home office ergonomics and equipment stipend upon completion of the probationary period. All remote home network connections must comply with IT Security Regulation SR-402.", body_style),
        ]
        doc.build(story)
    created_files.append(doc2026_pdf)

    # 3. Enterprise Security Regulation SR-402 PDF
    doc_sr402_pdf = os.path.join(target_dir, "Enterprise_Security_Regulation_SR402.pdf")
    if not os.path.exists(doc_sr402_pdf):
        doc = SimpleDocTemplate(doc_sr402_pdf, pagesize=letter)
        story = [
            Paragraph("Enterprise Security Regulation SR-402 (Cybersecurity & Access Compliance)", title_style),
            Paragraph("Document ID: REG-SR402 | Version: 3.1 | Effective Year: 2026 | Authority: Global InfoSec & Regulatory Compliance Board", body_style),
            Spacer(1, 15),
            Paragraph("Section 1. Regulatory Authority & Scope", h1_style),
            Paragraph("Regulation SR-402 establishes mandatory cybersecurity controls for all enterprise systems, data repositories, remote workstations, and employee processes (including Process X operations and Remote Workflows).", body_style),
            Spacer(1, 10),
            Paragraph("Section 2. Multi-Factor Authentication (MFA) & Encryption", h1_style),
            Paragraph("Requirement SEC-01: Multi-Factor Authentication (MFA) is strictly mandatory for all remote access and administrative tools. Hardware security keys or authenticator apps are required. SMS-based 2FA is prohibited. Requirement SEC-02: AES-256 encryption is mandated for all data at rest and TLS 1.3 for data in transit.", body_style),
            Spacer(1, 10),
            Paragraph("Section 3. Session Timeout & Inactivity Thresholds", h1_style),
            Paragraph("All remote workstation sessions and internal web application logins must automatically lock after 15 minutes of user inactivity. Screen lock passwords must satisfy the 14-character complexity matrix.", body_style),
            PageBreak(),
            Paragraph("Section 4. Data Retention & Audit Logging", h1_style),
            Paragraph("Requirement SEC-03: Security audit logs, user authentication traces, and system change records must be preserved for a minimum of 7 years in immutable cloud storage.", body_style),
            Spacer(1, 10),
            Paragraph("Section 5. Incident Reporting & Penalties", h1_style),
            Paragraph("Requirement SEC-04: Any suspected security breach or unauthorized data disclosure must be reported to the InfoSec Incident Response Team within 2 hours of detection. Failure to report within the 2-hour SLA carries Tier-3 disciplinary penalties and regulatory audit escalation.", body_style),
        ]
        doc.build(story)
    created_files.append(doc_sr402_pdf)

    # 4. IT Infrastructure Manual DOCX
    doc_it_path = os.path.join(target_dir, "IT_Infrastructure_Manual_v3.docx")
    if not os.path.exists(doc_it_path):
        doc_it = Document()
        doc_it.add_heading("IT Infrastructure Manual v3.4", 0)
        doc_it.add_paragraph("Document ID: MAN-IT-3.4 | Version: 3.4 | Department: Information Technology | Year: 2026")
        
        doc_it.add_heading("Section 1: Infrastructure Overview & Architecture", level=1)
        doc_it.add_paragraph("NexusCorp operates a hybrid cloud infrastructure across AWS and on-premise data centers. All production services adhere to zero-trust architecture principles governed by Regulation SR-402.")

        doc_it.add_heading("Section 2: Incident Response Workflow (Process X)", level=1)
        doc_it.add_paragraph("Process X defines the mission-critical incident triage and server provisioning pipeline for high-availability systems.")
        doc_it.add_paragraph("Step 1: Level 1 triage must be acknowledged within 15 minutes of automated alert triggering.")
        doc_it.add_paragraph("Step 2: If the incident severity is classified as High (Sev-1 or Sev-2), escalation to the Level 2 Engineering Team must occur within 1 hour.")
        doc_it.add_paragraph("Step 3: Process X requires full audit trace logging and compliance with Regulation SR-402 Section 5 incident notification window (2 hours maximum).")
        
        doc_it.add_heading("Section 3: Remote Access VPN Provisioning", level=1)
        doc_it.add_paragraph("Employees approved for remote work under Employee Operations Policy 2026 must be assigned an encrypted WireGuard/OpenVPN profile with split-tunneling disabled for corporate networks.")

        doc_it.save(doc_it_path)
    created_files.append(doc_it_path)

    # 5. Compliance Audit Guidelines XLSX
    doc_audit_path = os.path.join(target_dir, "Compliance_Audit_Guidelines_2026.xlsx")
    if not os.path.exists(doc_audit_path):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Audit_Checklist"
        ws.append(["Audit_ID", "Department", "Regulatory_Requirement", "Applicable_Policy", "Audit_Frequency", "Penalty_Level", "Status"])
        ws.append(["AUD-101", "Operations", "Attendance & Hybrid Schedule Review (60% on-site)", "POL-2026-V2 Section 2", "Biannual", "Level 1 (Internal Review)", "Compliant"])
        ws.append(["AUD-102", "Operations", "Remote Work Director Approvals", "POL-2026-V2 Section 3", "Quarterly", "Level 2 (HR Formal Notice)", "Compliant"])
        ws.append(["AUD-103", "Information Technology", "Process X Incident Escalation SLA (15m/1h)", "MAN-IT-3.4 Section 2", "Quarterly", "Level 2 (SLA Remediation)", "Compliant"])
        ws.append(["AUD-104", "Security & Compliance", "MFA Hardware Key Enforcement", "REG-SR402 Section 2", "Continuous", "Level 3 (Regulatory Breach)", "Compliant"])
        ws.append(["AUD-105", "Security & Compliance", "7-Year Audit Log Retention Policy", "REG-SR402 Section 4", "Annual", "Level 3 (Regulatory Penalty)", "Compliant"])
        ws.append(["AUD-106", "Security & Compliance", "2-Hour Incident Reporting Mandate", "REG-SR402 Section 5", "Continuous", "Level 3 (Immediate Suspension)", "Compliant"])
        wb.save(doc_audit_path)
    created_files.append(doc_audit_path)

    return created_files
